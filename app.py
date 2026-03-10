# app.py - VERSÃO DEFINITIVA (POSTGRESQL)

import os
import re
import fitz
import docx
import psycopg2
import psycopg2.extras
import click
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, flash
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_bcrypt import Bcrypt
import locale

# Esta linha descobre o caminho absoluto para o diretório onde app.py está
basedir = os.path.abspath(os.path.dirname(__file__))
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    locale.setlocale(locale.LC_ALL, '')

# --- CONFIGURAÇÃO ---
app = Flask(__name__)

app.secret_key = os.environ.get('SECRET_KEY', 'chave-local-para-nao-quebrar-o-teste')
UPLOAD_FOLDER = os.path.join(basedir, 'uploads')
GENERATED_FOLDER = os.path.join(basedir, 'generated')
TEMPLATE_FOLDER = os.path.join(basedir, 'templates_docx')

bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login' 
login_manager.login_message = 'Por favor, faça login para acessar esta página.'
login_manager.login_message_category = 'info'
app.config.update(
    UPLOAD_FOLDER=UPLOAD_FOLDER,
    GENERATED_FOLDER=GENERATED_FOLDER,
    TEMPLATE_FOLDER=TEMPLATE_FOLDER
)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)
os.makedirs(TEMPLATE_FOLDER, exist_ok=True)

# --- WRAPPER POSTGRESQL (Mantém a lógica do SQLite funcionando) ---
class DBWrapper:
    def __init__(self, conn):
        self.conn = conn

    def execute(self, query, params=None):
        cur = self.conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        if params:
            cur.execute(query, params)
        else:
            cur.execute(query)
        return cur

    def executemany(self, query, params_list):
        cur = self.conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cur.executemany(query, params_list)
        return cur

    def commit(self):
        self.conn.commit()

    def close(self):
        self.conn.close()

def get_db():
    # Conecta ao PostgreSQL usando a variável que salvamos no Render
    DATABASE_URL = os.environ.get('DATABASE_URL')
    if not DATABASE_URL:
        raise RuntimeError("DATABASE_URL não configurada no ambiente.")
    
    conn = psycopg2.connect(DATABASE_URL)
    return DBWrapper(conn)

@login_manager.user_loader
def load_user(user_id):
    db = get_db()
    user_data = db.execute('SELECT * FROM "user" WHERE id = %s', (user_id,)).fetchone()
    db.close()
    if user_data:
        return User(user_data['id'], user_data['username'], user_data['password_hash'])
    return None

# --- LÓGICA DE SUBSTITUIÇÃO DE TEXTO ---
def replace_text_in_paragraph(paragraph, key, value):
    if key not in paragraph.text:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if key not in full_text:
        return

    run_texts = [run.text for run in paragraph.runs]
    found = False
    for i in range(len(run_texts)):
        if key in run_texts[i]:
            run_texts[i] = run_texts[i].replace(key, value)
            found = True
            break
        
        combined_text = "".join(run_texts[i:])
        if combined_text.startswith(key):
            temp_text = ""
            j = i
            while len(temp_text) < len(key) and j < len(run_texts):
                temp_text += run_texts[j]
                j += 1
            
            run_texts[i] = temp_text.replace(key, value, 1)
            for k in range(i + 1, j):
                run_texts[k] = ""
            found = True
            break

    if found:
        for i in range(len(paragraph.runs)):
            paragraph.runs[i].text = run_texts[i]

# --- MODELO DE USUÁRIO ---
class User(UserMixin):
    def __init__(self, id, username, password_hash):
        self.id = id
        self.username = username
        self.password_hash = password_hash

    def get_id(self):
        return str(self.id)

# --- LÓGICA PRINCIPAL ---
def processar_pdf(pdf_path):
    try:
        texto_extraido = ""
        with fitz.open(pdf_path) as doc:
            for page in doc:
                texto_extraido += page.get_text() + " "

        texto_limpo = re.sub(r'\s+', ' ', texto_extraido.replace('\n', ' '))
        dados_do_projeto = {}

        padrao_tipo_e_numero = r"(PROJETO DE LEI (ORDIN[ÁA]RIA|COMPLEMENTAR)|PROJETO DE RESOLUÇÃO|PROJETO DE DECRETO LEGISLATIVO|PROPOSTA DE EMENDA [ÁA] LEI ORG[ÂA]NICA MUNICIPAL)\s*(?:N[º'q9]|n[oº9]|ne)\s*(\d+)"
        padrao_data = r"(\d{1,2}\s+de\s+\w+\s+(?:de|oe)\s+(\d{4}))"
        padrao_ementa = r"\"\s*(Abre.*?Anual.*?)\s*\""

        match_numero_val = None
        match_ano_val = None

        if (match := re.search(padrao_tipo_e_numero, texto_limpo, re.IGNORECASE)):
            dados_do_projeto["TIPO_PROJETO"] = match.group(1).upper().strip()
            match_numero_val = match.group(3)

        if (match_data := re.search(padrao_data, texto_limpo, re.IGNORECASE)):
            dados_do_projeto["DATA_PROJETO"] = match_data.group(1).strip()
            match_ano_val = match_data.group(2)

        if match_numero_val and match_ano_val:
            dados_do_projeto["NUMERO_PROJETO"] = f"{match_numero_val.zfill(3)}/{match_ano_val}"

        if (match := re.search(padrao_ementa, texto_limpo, re.IGNORECASE | re.DOTALL)):
            ementa_limpa = match.group(1).strip().replace("Í", "i").replace("çá", "çã").replace("ôe", "õe")
            dados_do_projeto["EMENTA"] = f'"{ementa_limpa}"'

        return dados_do_projeto
    except Exception as e:
        print(f"Erro ao processar PDF: {e}")
        return {}

def gerar_docx_final(form_data, pdf_filename):
    meses_pt = {
        1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
        5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
        9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
    }
    
    arquivos_gerados = []
    db = get_db()

    comissoes_selecionadas = form_data.getlist('comissao_selecionada')

    tipo_projeto = form_data.get("tipo_projeto", "").upper()
    autoria = form_data.get("autoria", "").upper()
    numero_projeto = form_data.get("numero_projeto", "00-0000")

    prefixo = "DOC"
    if "PROJETO DE LEI ORDINARIA" in tipo_projeto:
        prefixo = "PLOC" if "CÂMARA" in autoria else "PLOE"
    elif "PROJETO DE LEI COMPLEMENTAR" in tipo_projeto:
        prefixo = "PLCC" if "CÂMARA" in autoria else "PLCE"
    elif "PROJETO DE RESOLUÇÃO" in tipo_projeto:
        prefixo = "PRES"
    elif "PROJETO DE DECRETO LEGISLATIVO" in tipo_projeto:
        prefixo = "PDLC"
    elif "PROPOSTA DE EMENDA" in tipo_projeto:
        prefixo = "PELO"

    numero_sem_zero = numero_projeto.split('/')[0].lstrip('0')
    ano = numero_projeto.split('/')[-1]
    numero_formatado = f"{numero_sem_zero}_{ano}"

    for sigla in comissoes_selecionadas:
        template_path = os.path.join(app.config['TEMPLATE_FOLDER'], f"template_{sigla.lower()}.docx")

        if not os.path.exists(template_path): 
            continue

        doc = docx.Document(template_path)
        comissao = db.execute('SELECT * FROM comissoes WHERE sigla = %s', (sigla,)).fetchone()
        membros = db.execute('SELECT * FROM membros WHERE comissao_id = %s', (comissao['id'],)).fetchall()

        relator_id = form_data.get(f'relator_{sigla}')
        if not relator_id:
            continue 

        relator = db.execute('SELECT * FROM membros WHERE id = %s', (relator_id,)).fetchone()
        if not relator:
            continue

        signatarios = [m for m in membros if m['id'] != relator['id']]
        data_parecer = datetime.strptime(form_data.get('data_parecer'), '%Y-%m-%d')

        contexto = {
            "{{TIPO_PROJETO}}": form_data.get("tipo_projeto"),
            "{{NUMERO_PROJETO}}": form_data.get("numero_projeto"),
            "{{DATA_PROJETO}}": form_data.get("data_projeto", "").upper(),
            "{{EMENTA}}": form_data.get("ementa"),
            "{{AUTORIA}}": form_data.get("autoria"),
            "{{DATA_PROTOCOLO}}": datetime.strptime(form_data.get("data_protocolo"), '%Y-%m-%d').strftime('%d/%m/%Y'),
            "{{REGIME_URGENCIA}}": ", EM REGIME DE URGÊNCIA" if 'regime_urgencia' in form_data else "",
            "{{TEXTO_APRESENTACAO}}": f" e apresentada como objeto de deliberação na sessão ordinária do dia {datetime.strptime(form_data.get('data_apresentacao'), '%Y-%m-%d').strftime('%d/%m/%Y')}" if 'incluir_apresentacao' in form_data and form_data.get('data_apresentacao') else ".",
            "{{NUMERO_PARECER}}": form_data.get(f'num_parecer_{sigla}'),
            "{{DATA_PARECER_EXTENSO}}": f"{data_parecer.day:02d} de {meses_pt[data_parecer.month]} de {data_parecer.year}",
            "{{NOME_DA_COMISSAO}}": comissao['nome'].upper(),
            "{{NOME_RELATOR}}": relator['nome'].upper(),
            "{{CARGO_RELATOR}}": relator['cargo'],
            "{{NOME_SIGNATARIO_1}}": signatarios[0]['nome'].upper() if len(signatarios) > 0 else "",
            "{{CARGO_SIGNATARIO_1}}": signatarios[0]['cargo'] if len(signatarios) > 0 else "",
            "{{NOME_SIGNATARIO_2}}": signatarios[1]['nome'].upper() if len(signatarios) > 1 else "",
            "{{CARGO_SIGNATARIO_2}}": signatarios[1]['cargo'] if len(signatarios) > 1 else "",
        }

        for p in doc.paragraphs:
            for key, value in contexto.items():
                replace_text_in_paragraph(p, key, str(value)) 

        for table in doc.tables:
             for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in contexto.items():
                            replace_text_in_paragraph(p, key, str(value))

        nome_saida = f"{prefixo} {numero_formatado} {sigla}.docx"
        caminho_saida = os.path.join(app.config['GENERATED_FOLDER'], nome_saida)
        doc.save(caminho_saida)
        arquivos_gerados.append(nome_saida)

        db.execute('INSERT INTO pareceres (pdf_name, docx_name, numero_projeto, data_geracao) VALUES (%s, %s, %s, %s)',
                   (pdf_filename, nome_saida, form_data.get('numero_projeto'), datetime.now().strftime("%d/%m/%Y %H:%M:%S")))
        db.commit()

    db.close() 
    return arquivos_gerados

# --- ROTAS ---
@app.route('/', methods=['GET'])
@login_required
def index():
    db = get_db()
    historico = db.execute('SELECT * FROM pareceres ORDER BY id DESC').fetchall()
    return render_template('index.html', historico=historico)

@app.route('/upload', methods=['POST'])
@login_required
def upload():
    file = request.files.get('file')
    if not file or file.filename == '':
        flash('Nenhum arquivo selecionado.')
        return redirect(url_for('index'))
    
    filename = file.filename
    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(pdf_path)
    
    dados_pdf = processar_pdf(pdf_path)
    
    db = get_db()
    comissoes = db.execute('SELECT * FROM comissoes').fetchall()
    membros = db.execute('SELECT * FROM membros').fetchall()
    
    return render_template('revisar.html', dados=dados_pdf, comissoes=comissoes, membros=membros, filename=filename)

@app.route('/gerar', methods=['POST'])
@login_required
def gerar():
    comissoes_selecionadas = request.form.getlist('comissao_selecionada')
    if not comissoes_selecionadas:
        flash('Erro: Nenhuma comissão foi selecionada. Tente novamente.')
        return redirect(url_for('index'))

    pdf_filename = request.form.get('pdf_filename')
    
    try:
        arquivos_gerados = gerar_docx_final(request.form, pdf_filename)
        
        if not arquivos_gerados:
            flash('Erro ao gerar os arquivos. Verifique os templates e dados do formulário.')
            return redirect(url_for('index'))
            
        return render_template('resultado.html', arquivos=arquivos_gerados)

    except Exception as e:
        flash(f'Erro interno ao gerar documentos: {e}')
        return redirect(url_for('index'))

@app.route('/download/<filename>')
@login_required
def download(filename):
    return send_from_directory(app.config['GENERATED_FOLDER'], filename, as_attachment=True)

@app.route('/deletar_historico/<int:item_id>', methods=['POST'])
@login_required
def deletar_historico(item_id):
    try:
        db = get_db()
        item = db.execute('SELECT docx_name FROM pareceres WHERE id = %s', (item_id,)).fetchone()
        if item:
            arquivo_path = os.path.join(app.config['GENERATED_FOLDER'], item['docx_name'])
            if os.path.exists(arquivo_path):
                os.remove(arquivo_path)

        db.execute('DELETE FROM pareceres WHERE id = %s', (item_id,))
        db.commit()
        db.close()
        flash('Item do histórico removido com sucesso.', 'success')
    except Exception as e:
        flash(f'Erro ao remover item: {e}', 'danger')
    return redirect(url_for('index'))

@app.route('/limpar_historico', methods=['POST'])
@login_required
def limpar_historico():
    try:
        db = get_db()
        items = db.execute('SELECT docx_name FROM pareceres').fetchall()
        for item in items:
            arquivo_path = os.path.join(app.config['GENERATED_FOLDER'], item['docx_name'])
            if os.path.exists(arquivo_path):
                os.remove(arquivo_path)

        db.execute('DELETE FROM pareceres')
        db.commit()
        db.close()
        flash('Histórico completo removido com sucesso.', 'success')
    except Exception as e:
        flash(f'Erro ao limpar histórico: {e}', 'danger')
    return redirect(url_for('index'))

@app.route('/adicionar_membro', methods=['POST'])
@login_required
def adicionar_membro():
    try:
        nome = request.form['nome']
        cargo = request.form['cargo']
        comissao_id = request.form['comissao_id']
        
        db = get_db()
        db.execute('INSERT INTO membros (nome, cargo, comissao_id) VALUES (%s, %s, %s)',
                   (nome, cargo, comissao_id))
        db.commit()
        db.close()
        flash(f'Membro "{nome}" adicionado com sucesso!')
    except Exception as e:
        flash(f'Erro ao adicionar membro: {e}')
    return redirect(url_for('gerenciar'))

@app.route('/deletar_membro', methods=['POST'])
@login_required
def deletar_membro():
    try:
        membro_id = request.form['membro_id']
        
        db = get_db()
        nome_membro = db.execute('SELECT nome FROM membros WHERE id = %s', (membro_id,)).fetchone()['nome']
        db.execute('DELETE FROM membros WHERE id = %s', (membro_id,))
        db.commit()
        db.close()
        flash(f'Membro "{nome_membro}" removido com sucesso!')
    except Exception as e:
        flash(f'Erro ao remover membro: {e}')
    return redirect(url_for('gerenciar'))

@app.route('/editar_membro/<int:membro_id>', methods=['GET'])
@login_required
def editar_membro(membro_id):
    db = get_db()
    membro = db.execute('SELECT * FROM membros WHERE id = %s', (membro_id,)).fetchone()
    comissoes = db.execute('SELECT * FROM comissoes ORDER BY nome').fetchall()
    db.close()
    
    if membro is None:
        flash('Membro não encontrado.')
        return redirect(url_for('gerenciar'))
        
    return render_template('editar_membro.html', membro=membro, comissoes=comissoes)

@app.route('/atualizar_membro', methods=['POST'])
@login_required
def atualizar_membro():
    try:
        membro_id = request.form['membro_id']
        nome = request.form['nome']
        cargo = request.form['cargo']
        comissao_id = request.form['comissao_id']
        
        db = get_db()
        db.execute('UPDATE membros SET nome = %s, cargo = %s, comissao_id = %s WHERE id = %s',
                   (nome, cargo, comissao_id, membro_id))
        db.commit()
        db.close()
        flash(f'Dados do membro "{nome}" atualizados com sucesso!')
    except Exception as e:
        flash(f'Erro ao atualizar membro: {e}')
        
    return redirect(url_for('gerenciar'))

@app.route('/gerenciar')
@login_required
def gerenciar():
    db = get_db()
    comissoes = db.execute('SELECT * FROM comissoes ORDER BY nome').fetchall()
    
    membros_por_comissao = {}
    for comissao in comissoes:
        membros = db.execute(
            'SELECT * FROM membros WHERE comissao_id = %s ORDER BY nome', 
            (comissao['id'],)
        ).fetchall()
        membros_por_comissao[comissao['id']] = membros
    
    db.close()
    return render_template(
        'gerenciar.html', 
        comissoes=comissoes, 
        membros_por_comissao=membros_por_comissao
    )

@app.cli.command('create-admin')
@click.argument('username')
@click.argument('password')
def create_admin_command(username, password):
    db = get_db()
    try:
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        db.execute('INSERT INTO "user" (username, password_hash) VALUES (%s, %s)', (username, hashed_password))
        db.commit()
        print(f"Administrador '{username}' criado com sucesso.")
    except psycopg2.IntegrityError:
        db.conn.rollback()
        print(f"Erro: Usuário '{username}' já existe.")
    except Exception as e:
        print(f"Erro ao criar administrador: {e}")
    finally:
        db.close()

def init_db():
    db = get_db()

    print("Limpando tabelas antigas (se existiam)...")
    db.execute('DROP TABLE IF EXISTS pareceres CASCADE;')
    db.execute('DROP TABLE IF EXISTS membros CASCADE;')
    db.execute('DROP TABLE IF EXISTS comissoes CASCADE;')
    db.execute('DROP TABLE IF EXISTS "user" CASCADE;')

    print("Criando novas tabelas...")
    db.execute('''
    CREATE TABLE comissoes (
        id SERIAL PRIMARY KEY,
        nome TEXT NOT NULL,
        sigla TEXT NOT NULL UNIQUE
    );
    ''')

    db.execute('''
    CREATE TABLE membros (
        id SERIAL PRIMARY KEY,
        comissao_id INTEGER NOT NULL,
        nome TEXT NOT NULL,
        cargo TEXT NOT NULL,
        FOREIGN KEY (comissao_id) REFERENCES comissoes (id)
    );
    ''')

    db.execute('''
    CREATE TABLE pareceres (
        id SERIAL PRIMARY KEY,
        pdf_name TEXT NOT NULL,
        docx_name TEXT NOT NULL,
        numero_projeto TEXT,
        data_geracao TEXT NOT NULL
    );
    ''')

    db.execute('''
    CREATE TABLE "user" (
        id SERIAL PRIMARY KEY,
        username TEXT NOT NULL UNIQUE,
        password_hash TEXT NOT NULL
    );
    ''')

    comissoes = [
        ('Comissão de Justiça e Redação', 'CJR'),
        ('Comissão de Finanças e Orçamento', 'CFO'),
        ('Comissão de Obras, Serviços Públicos e Atividades Privadas', 'COSPAP'),
        ('Comissão de Educação, Saúde e Assistência Social', 'CESAS')
    ]
    db.executemany('INSERT INTO comissoes (nome, sigla) VALUES (%s, %s)', comissoes)
    db.commit()

    try:
        cjr_id = db.execute("SELECT id FROM comissoes WHERE sigla = 'CJR'").fetchone()['id']
        cfo_id = db.execute("SELECT id FROM comissoes WHERE sigla = 'CFO'").fetchone()['id']
        cospap_id = db.execute("SELECT id FROM comissoes WHERE sigla = 'COSPAP'").fetchone()['id']
        cesas_id = db.execute("SELECT id FROM comissoes WHERE sigla = 'CESAS'").fetchone()['id']
    except TypeError:
        print("ERRO: Falha ao buscar IDs das comissões. Verifique as siglas.")
        db.close()
        return

    membros = [
        (cjr_id, 'Vereador A (CJR)', 'Presidente'),
        (cjr_id, 'Vereador B (CJR)', 'Vice-Presidente'),
        (cjr_id, 'Vereador C (CJR)', 'Membro'),
        (cfo_id, 'Vereador D (CFO)', 'Presidente'),
        (cfo_id, 'Vereador E (CFO)', 'Vice-Presidente'),
        (cfo_id, 'Vereador F (CFO)', 'Membro'),
        (cospap_id, 'Vereador G (COSPAP)', 'Presidente'),
        (cospap_id, 'Vereador H (COSPAP)', 'Vice-Presidente'),
        (cospap_id, 'Vereador I (COSPAP)', 'Membro'),
        (cesas_id, 'Vereador J (CESAS)', 'Presidente'),
        (cesas_id, 'Vereador K (CESAS)', 'Vice-Presidente'),
        (cesas_id, 'Vereador L (CESAS)', 'Membro')
    ]
    db.executemany('INSERT INTO membros (comissao_id, nome, cargo) VALUES (%s, %s, %s)', membros)

    senha_admin = bcrypt.generate_password_hash("admin123").decode("utf-8")
    db.execute('INSERT INTO "user" (username, password_hash) VALUES (%s, %s) ON CONFLICT (username) DO NOTHING', ("admin", senha_admin))

    db.commit()
    db.close()
    print("Banco de dados PostgreSQL inicializado com sucesso!")

@app.cli.command('init-db')
def init_db_command():
    init_db()

@app.route('/setup-banco')
def setup_banco():
    try:
        init_db()
        return "<h3>Banco de dados PostgreSQL inicializado com sucesso!</h3><br><a href='/login'>Clique aqui para fazer o Login</a>"
    except Exception as e:
        return f"Erro ao criar o banco: {e}"

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
        
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        db = get_db()
        user_data = db.execute('SELECT * FROM "user" WHERE username = %s', (username,)).fetchone()
        db.close()
        
        if user_data:
            user = User(user_data['id'], user_data['username'], user_data['password_hash'])
            if bcrypt.check_password_hash(user.password_hash, password):
                login_user(user)
                flash('Login realizado com sucesso!', 'success')
                return redirect(url_for('index'))
                
        flash('Usuário ou senha inválidos.', 'danger')
        
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Você foi desconectado.', 'success')
    return redirect(url_for('login'))